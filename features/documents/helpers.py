from datetime import datetime, timedelta
from django.db.models import Q
from features.categories.models import Category
from .models import Document
import pymupdf
from docx import Document as DocxDocument
import unicodedata
import re
from .entities import Match, DeepSearchDocument
from dataclasses import asdict
from .serializers import DocumentSerializer

def get_descendants(category):
    children = [category]

    for child in category.childrens.all():
        children.extend(get_descendants(child))

    return children

def setDocumentFilters(data):
    filters = Q()
    category_qs = None

    if data.get("document_id"):
        filters &= Q(document_id__icontains=data["document_id"])

    if data.get("description"):
        filters &= Q(description__icontains=data["description"])

    if data.get("document_status"):
        if data.get("document_status") == "archived":
            filters &= Q(is_archived=True)
        elif data.get("document_status") == "recycled":
            filters &= Q(is_recycled=True)


    if data.get("document_name"):
        filters &= Q(document_name__icontains=data["document_name"])

    if data.get("staff_id"):
        filters &= Q(staff__id=data["staff_id"])

    if data.get("department_id"):
        filters &= Q(staff__department__id=data["department_id"])

    if data.get("dtype_id"):
        filters &= Q(dtype__id=data["dtype_id"])

    if data.get("dtype_name"):
        filters &= Q(dtype__name__icontains=data["dtype_name"])

    if data.get("created_at"):
        filters &= Q(created_at__date=data["created_at"])

    if data.get("updated_at"):
        filters &= Q(updated_at__date=data["updated_at"])

    if data.get("expired_at"):
        filters &= Q(expired_at=data["expired_at"])

    if data.get("from_date"):
        filters &= Q(created_at__gte=data["from_date"])

    if data.get("to_date"):
        to_date = datetime.strptime(data["to_date"], "%Y-%m-%d") + timedelta(days=1)
        filters &= Q(created_at__lt=to_date)

    if data.get("updated_from_date"):
        filters &= Q(updated_at__gte=data["updated_from_date"])

    if data.get("updated_to_date"):
        to_date = datetime.strptime(data["updated_to_date"], "%Y-%m-%d") + timedelta(days=1)
        filters &= Q(updated_at__lt=to_date)

    if data.get("expired_from_date"):
        filters &= Q(expired_at__gte=data["expired_from_date"])

    if data.get("expired_to_date"):
        to_date = datetime.strptime(data["expired_to_date"], "%Y-%m-%d") + timedelta(days=1)
        filters &= Q(expired_at__lt=to_date)

    if data.get("category_id"):
        category = Category.objects.get(id=data["category_id"])

        if not category.childrens.exists():
            filters &= Q(category__id=data["category_id"])
        else:
            category_qs = get_descendants(category)

    return filters, category_qs


def getDocumentName():
    document_names = Document.objects.values_list(Document.document_name)
    return document_names


def _normalize(text: str) -> str:
    """
    Normalize Unicode text for consistent comparison.
    - NFC normalization covers Burmese Unicode (handles composed vs decomposed forms).
    - Strip zero-width joiners, non-breaking spaces, and other invisible chars
      that commonly appear in Burmese PDF extraction.
    """
    text = unicodedata.normalize("NFC", text)
    text = re.sub(r"[\u200b\u200c\u200d\ufeff\u00ad]", "", text)
    return text


def _search_pdf(file_path, normalized_key):
    """Extract text from PDF and return matches with page and line numbers."""
    matches = []

    try:
        doc = pymupdf.open(file_path)
    except Exception as e:
        print(f"Failed to open PDF {file_path}: {e}")
        return matches

    for page_num, page in enumerate(doc, start=1):

        try:
            raw_text = page.get_text()
        except Exception as e:
            print(f"Failed to extract text from page {page_num}: {e}")
            continue

        lines = raw_text.splitlines()

        for line_num, line in enumerate(lines, start=1):

            normalized_line = _normalize(line).lower()

            if normalized_key in normalized_line:
                matches.append({
                    "text": line.strip(),
                    "line": line_num,
                    "page": page_num,
                })

    doc.close()
    return matches


def _search_docx(file_path, normalized_key):
    """Extract text from DOCX and return matches with line numbers only."""
    matches = []

    try:
        doc = DocxDocument(file_path)
    except Exception as e:
        print(f"Failed to open DOCX {file_path}: {e}")
        return matches

    for line_num, paragraph in enumerate(doc.paragraphs, start=1):

        line = paragraph.text

        if not line.strip():
            continue

        normalized_line = _normalize(line).lower()

        if normalized_key in normalized_line:
            matches.append({
                "text": line.strip(),
                "line": line_num,
                "page": None,
            })

    return matches


def deepSearch(documents, keys):

    if not keys:
        return []

    document_list = []

    normalized_key = _normalize(keys).lower()

    for document in documents:
        print("this is document name : ", document.document_name)

        file_path = document.document.path
        ext = file_path.rsplit(".", 1)[-1].lower() if "." in file_path else ""

        if ext == "pdf":
            file_matches = _search_pdf(file_path, normalized_key)
        elif ext in ("docx", "doc"):
            file_matches = _search_docx(file_path, normalized_key)
        else:
            print(f"Unsupported file type '{ext}' for {document.document_name}, skipping.")
            continue

        if file_matches:

            document_list.append(
                DeepSearchDocument(
                    id=document.id,
                    filename=document.document_name,
                    department_name=document.staff.department.department_name,
                    uploaded_at=document.created_at,
                    category=document.category.category_name,
                    file_url=document.document.path,
                    matches=[
                        Match(
                            page=m["page"],
                            line=m["line"],
                            text=m["text"],
                        )
                        for m in file_matches
                    ]
                )
            )

    return [asdict(item) for item in document_list]




